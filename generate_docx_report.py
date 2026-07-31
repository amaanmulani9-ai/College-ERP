import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def make_callout(doc, text, title="KEY ARCHITECTURAL HIGHLIGHT"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "EBF8FF")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border blue
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="3182CE"/>
            <w:top w:val="none"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    
    run_b = p.add_run(text)
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # Empty paragraph after callout for spacing
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(6)

def style_heading_1(p):
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(18)
        r.bold = True
        r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy

def style_heading_2(p):
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = RGBColor(0x2C, 0x52, 0x82) # Slate Blue

def style_heading_3(p):
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = RGBColor(0x2D, 0x37, 0x48) # Dark Slate

def add_body_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    return p

def add_bullet_p(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    r_bold = p.add_run(bold_prefix)
    r_bold.bold = True
    r_bold.font.name = "Calibri"
    r_bold.font.size = Pt(10.5)
    r_bold.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    r_text = p.add_run(text)
    r_text.font.name = "Calibri"
    r_text.font.size = Pt(10.5)
    r_text.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    return p

def format_table(doc, tbl, headers, rows_data, col_widths=None):
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    # Header Row
    hdr_cells = tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1B365D")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(10)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    # Data Rows
    for r_idx, row_values in enumerate(rows_data):
        row_cells = tbl.add_row().cells
        bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_values):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=90, bottom=90, left=120, right=120)
            p = row_cells[c_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
                
    # Column Widths
    if col_widths:
        for row in tbl.rows:
            for c_idx, w in enumerate(col_widths):
                row.cells[c_idx].width = Inches(w)
                
    set_table_borders(tbl, color="CBD5E0", sz="4")
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(8)

def main():
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    # Title Header Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(4)
    r_t = p_title.add_run("COLLEGE ERP SYSTEM")
    r_t.font.name = "Calibri"
    r_t.font.size = Pt(26)
    r_t.bold = True
    r_t.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(16)
    r_s = p_sub.add_run("Full System Architecture, Business Logic & Programming Language Specification Report")
    r_s.font.name = "Calibri"
    r_s.font.size = Pt(14)
    r_s.italic = True
    r_s.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(24)
    r_m = p_meta.add_run("Author: Amaan | Technology Stack: Python, Django, PostgreSQL, React, AI Suite | Date: July 2026")
    r_m.font.name = "Calibri"
    r_m.font.size = Pt(10)
    r_m.bold = True
    r_m.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # SECTION 1: EXECUTIVE SUMMARY & PROJECT SYNOPSIS
    # ----------------------------------------------------
    h1 = doc.add_paragraph("1. Executive Summary & Project Synopsis")
    style_heading_1(h1)
    
    add_body_p(doc, "The College ERP System is a next-generation, multi-tenant SaaS-enabled Enterprise Resource Planning platform tailored for higher educational institutions (universities, colleges, and polytechnics). The system bridges all institutional operations—bringing together Administration, Teaching Staff, Students, Parents, Alumni, Back-office Staff, and Corporate Recruiters under a single unified web and mobile interface.")
    
    make_callout(doc, "College ERP utilizes PostgreSQL schema-level isolation (via django-tenants) to ensure 100% data separation between colleges on a shared infrastructure, alongside an integrated Smart AI Suite for local LLM exam creation, predictive student risk analysis, and fee receipt fraud detection.", "MULTI-TENANT SAAS & SMART AI PLATFORM")
    
    add_body_p(doc, "The system caters to 7 core user personas, each accessing custom, role-bounded portals:")
    add_bullet_p(doc, "1. HOD / Administrator (Role 1): ", "Full institutional governance—managing departments, courses, batch allocations, staff appointments, fee policy structures, systemic notifications, leave approvals, and executive analytics.")
    add_bullet_p(doc, "2. Staff / Teacher (Role 2): ", "Academic execution—recording class attendance via manual or QR barcode terminals, entering internal/semester grades, hosting virtual classrooms, managing library book issues, and drafting assignments/questions.")
    add_bullet_p(doc, "3. Student (Role 3): ", "Academic self-service—viewing real-time attendance statistics, timetables, downloading study material, submitting assignments, paying fees via online gateway, generating AI resumes, borrowing library books, and attending virtual lectures.")
    add_bullet_p(doc, "4. Parent (Role 4): ", "Academic & financial oversight—tracking child's daily attendance records, exam performance trends, fee schedules, and communicating with administration.")
    add_bullet_p(doc, "5. Alumni (Role 5): ", "Post-graduation network—connecting with current campus recruitment drives, accessing career advancement services, and tracking institutional updates.")
    add_bullet_p(doc, "6. Company HR / Recruiter (Role 6): ", "Talent acquisition—posting job opportunities, searching student skill databases, scheduling interviews, and shortlisting candidates.")
    add_bullet_p(doc, "7. Back Office Staff (Role 7): ", "Administrative accounting—verifying applicant documents during admissions, managing physical document archives, executing fee collection counters, and issuing degree certificates.")

    # ----------------------------------------------------
    # SECTION 2: PROGRAMMING LANGUAGES & TECH STACK
    # ----------------------------------------------------
    h1 = doc.add_paragraph("2. Comprehensive Technology Stack & Languages Used")
    style_heading_1(h1)
    
    add_body_p(doc, "The College ERP project is constructed using a robust polyglot architecture, combining high-performance backend processing, dynamic web templates, mobile compilation targets, and specialized AI/ML libraries.")
    
    h2 = doc.add_paragraph("2.1 Programming & Scripting Languages")
    style_heading_2(h2)
    
    lang_headers = ["Language", "Usage Layer", "Primary Application & Purpose"]
    lang_rows = [
        ["Python 3.8 - 3.14+", "Backend Core & AI/ML Services", "Drives the main web server, business logic views, ORM data models, REST APIs, Celery async background tasks, machine learning predictive analytics, document OCR, and Ollama LLM integration."],
        ["JavaScript (ES6+)", "Frontend Client & Interactive Logic", "Powers client-side dynamic UI interactions, Chart.js rendering, HTML5 webcam QR scanning (HTML5-QRcode/Instascan), CODE128 barcode generation (JsBarcode), and real-time WebSocket chat client."],
        ["HTML5", "UI Structure & Templating", "Structures web layouts using Django HTML template tags, AdminLTE layouts, accessible forms, and multi-step registration wizards."],
        ["CSS3 / Vanilla CSS", "Styling & Responsive Design", "Provides visual styling, custom responsive flex/grid layouts, AdminLTE dashboard themes, liquid glass card components, and CSS animations without heavy frameworks."],
        ["SQL (PostgreSQL)", "Relational Database", "Executes relational queries, multi-schema DDL/DML partitioning, transactional fee records, and indexed search operations."],
        ["Kotlin / Java", "Mobile Client (Android APK)", "Powers the native Android mobile application wrapper (`CampusPro_College_ERP.apk`) for student and staff mobile devices."],
        ["PowerShell & Shell (Bash)", "Build & Automation Scripts", "Automates environment setup, dependency checks, database seeding (`seed_erp_data.py`), testing execution (`run_verification.ps1`), and APK compilation (`build_apk.ps1`)."],
        ["YAML & Dockerfile", "Containerization & Orchestration", "Configures Docker containers (`docker-compose.yml`), Kubernetes deployment manifests (`k8s/`), NGINX routing, and Vercel/Render cloud services."]
    ]
    
    tbl_lang = doc.add_table(rows=1, cols=3)
    format_table(doc, tbl_lang, lang_headers, lang_rows, col_widths=[1.5, 1.8, 3.2])
    
    h2 = doc.add_paragraph("2.2 Technology Stack Breakdown")
    style_heading_2(h2)
    
    add_bullet_p(doc, "Backend Framework: ", "Django 5.x with Django REST Framework (DRF) for API endpoints, `django-tenants` for multi-tenancy schema switching, Gunicorn as WSGI server, WhiteNoise for static file serving, and Pytest / Django Test Framework for unit & integration testing.")
    add_bullet_p(doc, "Database & Storage: ", "PostgreSQL 15 as primary relational database with schema-based isolation, Redis for memory caching, session storage, and Celery task broker, MinIO for S3-compatible asset and certificate document storage.")
    add_bullet_p(doc, "Frontend UI & Visualization: ", "Django HTML Template Engine, Custom Vanilla CSS, Bootstrap 4/5, AdminLTE Dashboard Skeleton, Chart.js for analytical charts, JsBarcode for printable student ID barcodes, HTML5-QRcode & Instascan for camera QR attendance, Jitsi Meet IFrame SDK for virtual classrooms.")
    add_bullet_p(doc, "Authentication & Security: ", "Keycloak OpenID Connect (OIDC) Single Sign-On (SSO), Django Custom User Auth with email login, CSRF tokens, Session Middleware in Redis, Firebase Cloud Messaging (FCM) for mobile push notifications, Cryptographic SHA-256 receipt verification.")
    add_bullet_p(doc, "Smart AI & Machine Learning Suite: ", "Ollama hosting local open-weight LLMs (Llama 3, Mistral) for exam paper and timetable generation, ChromaDB native vector database for RAG context retrieval, Scikit-learn based predictive analytics engine for student risk detection, OpenCV/Face Recognition service, Document OCR service.")
    add_bullet_p(doc, "Infrastructure & DevOps: ", "NGINX Reverse Proxy, Docker & Docker Compose, Kubernetes manifests, Meilisearch Engine for instant fuzzy search, Vercel & Supabase integration support, Prometheus & Grafana monitoring stack.")

    # ----------------------------------------------------
    # SECTION 3: SYSTEM ARCHITECTURE & MULTI-TENANCY
    # ----------------------------------------------------
    h1 = doc.add_paragraph("3. System Architecture & Multi-Tenancy Engine")
    style_heading_1(h1)
    
    add_body_p(doc, "The system architecture is engineered for multi-tenant SaaS scalability. Multiple educational institutions run on a single shared codebase and database server while maintaining 100% data privacy and security.")
    
    h2 = doc.add_paragraph("3.1 Shared-Model Schema Isolation (`django-tenants`)")
    style_heading_2(h2)
    
    add_body_p(doc, "Rather than maintaining separate database servers or filtering every query by a `college_id` foreign key (which is error-prone and risks data leaks), College ERP uses PostgreSQL Schema Isolation:")
    
    make_callout(doc, "PostgreSQL Database Structure:\n• Public Schema ('public'): Contains tenant definitions (saas_admin_client) and domain mappings (saas_admin_domain).\n• Tenant Schema A ('college1_db'): Isolated tables for Student, Staff, Attendance, Fees, Exams.\n• Tenant Schema B ('college2_db'): Completely separate tables for another institution.", "SCHEMA-LEVEL MULTI-TENANCY")
    
    add_body_p(doc, "When an HTTP request arrives, the custom `TenantMainMiddleware` intercepts the incoming request hostname (e.g. `stanford.college-erp.com`), looks up the domain in the `public` schema domain registry, and dynamically executes a SQL command to switch the active database search path:")
    
    add_body_p(doc, "SET search_path TO stanford_db, public;")
    
    h2 = doc.add_paragraph("3.2 HTTP Request Middleware Lifecycle Pipeline")
    style_heading_2(h2)
    
    add_body_p(doc, "Every incoming request passes sequentially through the following middleware layers before reaching the view logic:")
    add_bullet_p(doc, "1. TenantMainMiddleware: ", "Extracts hostname and sets active PostgreSQL DB tenant schema search path.")
    add_bullet_p(doc, "2. SecurityMiddleware: ", "Enforces SSL HTTPS redirects and sets security headers (X-Frame-Options, X-Content-Type-Options).")
    add_bullet_p(doc, "3. SessionMiddleware: ", "Retrieves and validates user session state from Redis.")
    add_bullet_p(doc, "4. LocaleMiddleware: ", "Detects user language preferences for internationalization (i18n).")
    add_bullet_p(doc, "5. CsrfViewMiddleware: ", "Verifies CSRF tokens on POST/PUT requests to prevent Cross-Site Request Forgery.")
    add_bullet_p(doc, "6. AuthenticationMiddleware: ", "Authenticates user session or Keycloak OIDC JWT token, populating `request.user`.")
    add_bullet_p(doc, "7. MessageMiddleware: ", "Manages flash notification messages (e.g., success, error popups).")
    add_bullet_p(doc, "8. WhiteNoiseMiddleware: ", "Serves compressed static CSS/JS assets directly from python server.")

    # ----------------------------------------------------
    # SECTION 4: DATABASE SCHEMA & DATA MODELS ANALYSIS
    # ----------------------------------------------------
    h1 = doc.add_paragraph("4. Database Schema & Data Models Breakdown")
    style_heading_1(h1)
    
    add_body_p(doc, "All tenant operational data is managed via Django ORM models defined in `backend/main_app/models.py`. The core models and their relationships are structured as follows:")
    
    model_headers = ["Model Name", "Key Attributes & Fields", "Business Logic & Relationships"]
    model_rows = [
        ["CustomUser", "id, email, password, user_type (1-7), first_name, last_name, profile_pic, fcm_token, created_at", "Extends `AbstractUser`. Uses `email` as username. `user_type` determines access role (HOD=1, Staff=2, Student=3, Parent=4, Alumni=5, HR=6, Backoffice=7). Links 1-to-1 to Admin, Staff, Student, Backoffice profiles."],
        ["Student", "id, admin (FK User), gender, address, course_id (FK), session_year_id (FK), unique_student_code, id_card_code, batch_year", "Student profile. Generates unique student registration code and UUID-based barcode/QR string (`id_card_code`) automatically upon saving."],
        ["Staff", "id, admin (FK User), address, mobile_number, department, joining_date, salary", "Faculty profile. Linked to department and assigned subjects. Used for attendance entry and exam paper creation."],
        ["AdminHOD", "id, admin (FK User), created_at, updated_at", "Administrative HOD profile overseeing campus operations."],
        ["Course", "id, course_name, monthly_fees, total_semesters, created_at", "Degree/Diploma program definition (e.g., B.Tech Computer Engineering). Stores fee structures and semester limits."],
        ["Subject", "id, subject_name, subject_code, course_id (FK), staff_id (FK User)", "Individual subject module. Linked to a parent Course and taught by an assigned Staff member."],
        ["SessionYearModel", "id, session_start_year, session_end_year", "Academic academic year (e.g., 2024-2028). Used to partition student cohorts."],
        ["Attendance", "id, subject_id (FK), attendance_date, session_year_id (FK)", "Daily attendance session record for a subject on a specific date."],
        ["AttendanceReport", "id, student_id (FK), attendance_id (FK), status (Boolean)", "Individual student attendance status (True=Present, False=Absent) linked to an Attendance session."],
        ["StudentResult", "id, student_id (FK), subject_id (FK), subject_exam_marks, subject_assignment_marks", "Academic marks entry storing test scores and assignment grades."],
        ["FeeRecord", "id, student_id (FK), title, total_amount, paid_amount, due_date, status", "Billed invoice for student tuition, exam, or hostel fees."],
        ["FeePayment", "id, fee_record (FK), payment_amount, transaction_id, payment_mode, receipt_code", "Fee settlement transaction. Generates cryptographic receipt verification codes."],
        ["PlacementDrive", "id, company_name, job_title, salary_package, drive_date, min_cgpa", "Campus recruitment drive listing posted by Company HR."],
        ["PlacementRegistration", "id, placement_drive (FK), student (FK), status, resume_url", "Student registration for a specific campus recruitment drive."],
        ["LiveClass", "id, subject (FK), staff (FK), room_name, scheduled_at, is_active", "Virtual classroom session integrating Jitsi Meet video conferencing."],
        ["IssuedBook", "id, student (FK), book_isbn, book_title, issue_date, return_date, is_returned", "Library book lending transaction log."]
    ]
    
    tbl_models = doc.add_table(rows=1, cols=3)
    format_table(doc, tbl_models, model_headers, model_rows, col_widths=[1.5, 2.2, 2.8])

    # ----------------------------------------------------
    # SECTION 5: DETAILED BUSINESS LOGIC & MODULE WORKFLOWS
    # ----------------------------------------------------
    h1 = doc.add_paragraph("5. Detailed Business Logic & Module Workflows")
    style_heading_1(h1)
    
    add_body_p(doc, "The business logic of College ERP is decoupled across specialized Django view modules inside `backend/main_app/`. Below is a breakdown of how key institutional workflows operate:")
    
    h2 = doc.add_paragraph("5.1 HOD & Administrative Governance (`hod_views.py`)")
    style_heading_2(h2)
    add_body_p(doc, "The HOD module contains over 130KB of administrative business logic:")
    add_bullet_p(doc, "• Staff & Student Lifecycle: ", "Handles onboarding, profile edits, course assignments, and deletion with cascading constraints.")
    add_bullet_p(doc, "• Attendance Auditing: ", "Provides administrative override capabilities to modify incorrect attendance logs recorded by staff.")
    add_bullet_p(doc, "• Leave Management: ", "Reviews leave applications submitted by staff and students, updating approval statuses (`1=Approved, 2=Disapproved`).")
    add_bullet_p(doc, "• Institutional Analytics: ", "Calculates department-wide pass percentages, attendance averages, and fee collection totals.")

    h2 = doc.add_paragraph("5.2 Student Digital ID & QR Attendance System (`smart_views.py`)")
    style_heading_2(h2)
    add_body_p(doc, "Attendance management supports both traditional manual roll-call and high-speed QR barcode scanning:")
    add_bullet_p(doc, "1. ", "Every student is assigned a unique `id_card_code` (UUID token) rendered on their printable digital ID card using JsBarcode (CODE128) and HTML5 QR generators.")
    add_bullet_p(doc, "2. ", "During class, the teacher opens the QR Scanner Terminal (`/staff/attendance/take/`).")
    add_bullet_p(doc, "3. ", "When the student presents their ID card, the camera feed (powered by HTML5-QRcode) scans the code and posts it to `smart_views.py`.")
    add_bullet_p(doc, "4. ", "The backend verifies the student's enrollment, checks if an `Attendance` object exists for today's date/subject, creates it if necessary, and inserts/updates an `AttendanceReport` record with `status=True`.")

    h2 = doc.add_paragraph("5.3 Smart AI Suite & Exam Paper Generator (`ai_views.py` & `ai_helper.py`)")
    style_heading_2(h2)
    add_body_p(doc, "The Smart AI Suite embeds artificial intelligence directly into academic routines:")
    add_bullet_p(doc, "• Ollama Local LLM Inference: ", "Teachers select a Course and Subject and provide a prompt (e.g. 'Generate 10 multiple-choice questions on Data Structures Binary Trees with options and answer keys'). The backend sends context to an Ollama Docker container running Llama 3 or Mistral, parses the returned JSON structure, and automatically creates draft exam questions.")
    add_bullet_p(doc, "• AI Resume Builder: ", "Students input project details and skills into `ai_helper.py`, which formats and generates optimized ATS-friendly resumes for placement drives.")
    add_bullet_p(doc, "• Fee Receipt Fraud Detection: ", "Cryptographically signs fee receipts with SHA-256 hashes containing student code, payment timestamp, and transaction ID to detect tampered payment receipts.")

    h2 = doc.add_paragraph("5.4 Financial Management & Razorpay Gateway (`finance_views.py`)")
    style_heading_2(h2)
    add_body_p(doc, "Financial logic manages student billing cycles:")
    add_bullet_p(doc, "1. ", "HOD creates `FeeRecord` items for batches (e.g. Semester 4 Tuition Fee - $1,200).")
    add_bullet_p(doc, "2. ", "Students log into `/student/payable-fees/` and view outstanding balances.")
    add_bullet_p(doc, "3. ", "Clicking 'Pay Now' initializes a Razorpay Order API call, generating a payment transaction token.")
    add_bullet_p(doc, "4. ", "Upon successful payment, Razorpay webhook triggers `FeePayment` creation, marks `FeeRecord.status = Paid`, and generates an official PDF fee receipt with QR verification.")

    h2 = doc.add_paragraph("5.5 Predictive Student Analytics Engine (`analytics_engine.py`)")
    style_heading_2(h2)
    add_body_p(doc, "The analytics engine aggregates student data across attendance, internal test marks, assignment submission frequency, and library activity to calculate an 'Academic At-Risk Score':")
    make_callout(doc, "If Attendance < 75% AND Exam Marks < 40%, Risk Level = HIGH (Red Flagged).\nAutomated alerts are dispatched to the Student's HOD and Parent portal for early faculty intervention.", "PREDICTIVE STUDENT RISK ANALYSIS")

    # ----------------------------------------------------
    # SECTION 6: INTERACTIVE DATA FLOWS & DIAGRAMS
    # ----------------------------------------------------
    h1 = doc.add_paragraph("6. Interactive Data Flows & Sequence Diagrams")
    style_heading_1(h1)
    
    add_body_p(doc, "Below are detailed step-by-step sequence flows illustrating core logic execution:")
    
    h2 = doc.add_paragraph("6.1 Student QR Attendance Log Flow")
    style_heading_2(h2)
    
    flow1_headers = ["Step #", "Actor / Component", "Action & Logic Executed", "Destination / Output"]
    flow1_rows = [
        ["1", "Student", "Presents Digital ID Card (containing QR code with `id_card_code`)", "Teacher / Scanner Terminal"],
        ["2", "Teacher / Web UI", "Scans QR code using webcam feed powered by HTML5-QRcode", "HTTP POST to `smart_views.py`"],
        ["3", "Backend (`smart_views`)", "Queries `Student.objects.get(id_card_code=code)`", "PostgreSQL Tenant Schema"],
        ["4", "Database", "Returns Student record, active Batch, and enrolled Subject", "Backend"],
        ["5", "Backend", "Checks if `Attendance(date=today, subject=subject)` exists; creates if missing", "Attendance Instance"],
        ["6", "Backend", "Executes `AttendanceReport.objects.update_or_create(student=student, status=True)`", "AttendanceReport Record"],
        ["7", "Scanner Terminal", "Displays green notification: 'Attendance Logged: [Student Name] - Present'", "Teacher Dashboard"]
    ]
    
    tbl_flow1 = doc.add_table(rows=1, cols=4)
    format_table(doc, tbl_flow1, flow1_headers, flow1_rows, col_widths=[0.6, 1.5, 2.6, 1.8])

    h2 = doc.add_paragraph("6.2 AI Exam Paper Generation Flow")
    style_heading_2(h2)
    
    flow2_headers = ["Step #", "Actor / Component", "Action & Logic Executed", "Destination / Output"]
    flow2_rows = [
        ["1", "Teacher", "Selects Course, Subject, difficulty, and question count prompt", "AI Suite UI (`ai_views.py`)"],
        ["2", "Backend (`ai_views`)", "Fetches Subject syllabus keywords & historical exam papers", "PostgreSQL Database"],
        ["3", "Backend", "Formats prompt with JSON schema constraints and sends POST request", "Ollama LLM Container (Port 11434)"],
        ["4", "Ollama (Llama 3)", "Runs local LLM inference to draft questions, choices, and answer key", "Returns structured JSON payload"],
        ["5", "Backend", "Parses JSON payload and executes `Question.objects.bulk_create()`", "Saved to PostgreSQL Database"],
        ["6", "Teacher Dashboard", "Renders formatted question paper draft for teacher review and editing", "Web Browser View"]
    ]
    
    tbl_flow2 = doc.add_table(rows=1, cols=4)
    format_table(doc, tbl_flow2, flow2_headers, flow2_rows, col_widths=[0.6, 1.5, 2.6, 1.8])

    # ----------------------------------------------------
    # SECTION 7: FILE BASE & DIRECTORY MAPPING
    # ----------------------------------------------------
    h1 = doc.add_paragraph("7. Codebase Directory Structure & File Mapping")
    style_heading_1(h1)
    
    add_body_p(doc, "The repository is structured into modular backend components and template layouts:")
    
    file_headers = ["File Path", "Size / Category", "Key Responsibility & Logic Contained"]
    file_rows = [
        ["backend/main_app/models.py", "61.7 KB", "Defines all ORM data models (CustomUser, Student, Staff, Attendance, FeeRecord, PlacementDrive, LiveClass, etc.)."],
        ["backend/main_app/hod_views.py", "139.3 KB", "Core HOD administrative views (staff/student management, course setups, leave approvals, analytics)."],
        ["backend/main_app/student_views.py", "71.8 KB", "Student portal logic (attendance viewing, timetable, fees, report cards, library books, resume builder)."],
        ["backend/main_app/views.py", "58.2 KB", "Authentication, user login, logout, password resets, profile management, registration forms."],
        ["backend/main_app/staff_views.py", "32.2 KB", "Teacher portal logic (manual/QR attendance entry, results submission, assignment upload, virtual class scheduling)."],
        ["backend/main_app/chat_views.py", "18.3 KB", "Real-time chat, student-teacher discussion boards, message routing."],
        ["backend/main_app/ai_views.py & ai_helper.py", "29.6 KB", "AI exam question generator, Ollama API connector, RAG vector lookup helper, resume generator."],
        ["backend/main_app/mobile_api_views.py", "15.0 KB", "REST API endpoints for Android APK (`CampusPro`) and mobile PWA sync."],
        ["backend/main_app/backoffice_views.py", "13.7 KB", "Back-office operations (document verification, admissions processing, degree certificates)."],
        ["backend/main_app/analytics_views.py", "13.1 KB", "Data analytics endpoints powering Chart.js dashboard widgets and predictive risk charts."],
        ["backend/main_app/smart_views.py", "11.0 KB", "QR code/barcode attendance scanning logic, digital ID validation."],
        ["backend/main_app/finance_views.py", "10.3 KB", "Fee billing calculations, payment receipts, Razorpay webhook handling."],
        ["frontend/templates/main_app/", "Templates", "Base layout structures (`base.html`), login page, navigation sidebars (`erpnext_sidebar.html`)."],
        ["frontend/templates/hod_template/", "Templates", "HOD admin dashboard templates, management panels, batch setups."],
        ["frontend/templates/staff_template/", "Templates", "Teacher attendance desks, marks entry forms, assignment managers."],
        ["frontend/templates/student_template/", "Templates", "Student home, attendance view, timetable, fee payment desk, book library."]
    ]
    
    tbl_files = doc.add_table(rows=1, cols=3)
    format_table(doc, tbl_files, file_headers, file_rows, col_widths=[2.2, 1.2, 3.1])

    # ----------------------------------------------------
    # SECTION 8: PAGE & ROUTE AUDIT
    # ----------------------------------------------------
    h1 = doc.add_paragraph("8. Registered Routes & Page Status Audit")
    style_heading_1(h1)
    
    add_body_p(doc, "A comprehensive backend audit comparing registered URL routes (`urls.py`) against rendering template files reveals the operational status of all platform pages:")
    
    h2 = doc.add_paragraph("8.1 Live & Verified Pages")
    style_heading_2(h2)
    
    audit_headers = ["URL Route Path", "View Handler", "Template Rendered", "Role Access", "Audit Status"]
    audit_rows = [
        ["/", "views.login_page", "main_app/login.html", "Public / All", "LIVE & VERIFIED"],
        ["/admin/home/", "hod_views.admin_home", "hod_template/home_content.html", "HOD / Admin", "LIVE & VERIFIED"],
        ["/admin/batches/", "hod_views.admin_manage_batches", "hod_template/manage_batches.html", "HOD / Admin", "LIVE & VERIFIED"],
        ["/staff/home/", "staff_views.staff_home", "staff_template/erpnext_staff_home.html", "Staff / Teacher", "LIVE & VERIFIED"],
        ["/staff/attendance/take/", "staff_views.staff_take_attendance", "staff_template/staff_take_attendance.html", "Staff / Teacher", "LIVE & VERIFIED"],
        ["/student/home/", "student_views.student_home", "student_template/erpnext_student_home.html", "Student", "LIVE & VERIFIED"],
        ["/student/timetable/", "student_views.student_timetable", "student_template/student_timetable.html", "Student", "LIVE & VERIFIED"],
        ["/student/payable-fees/", "student_views.student_payable_fees", "student_template/student_payable_fees.html", "Student", "LIVE & VERIFIED"],
        ["/student/viewbooks/", "student_views.view_books", "student_template/view_books.html", "Student", "LIVE & VERIFIED"],
        ["/student/report-card/", "student_views.student_report_card", "student_template/student_report_card.html", "Student", "LIVE & VERIFIED"],
        ["/chat/", "chat_views.chat_home", "main_app/chat.html", "All Authenticated", "LIVE & VERIFIED"]
    ]
    
    tbl_audit = doc.add_table(rows=1, cols=5)
    format_table(doc, tbl_audit, audit_headers, audit_rows, col_widths=[1.5, 1.5, 1.8, 0.9, 0.8])

    # ----------------------------------------------------
    # SECTION 9: DEPLOYMENT & VERIFICATION GUIDE
    # ----------------------------------------------------
    h1 = doc.add_paragraph("9. Local Execution & Automated Verification Guide")
    style_heading_1(h1)
    
    add_body_p(doc, "To setup and run the College ERP system locally for development or testing:")
    
    add_bullet_p(doc, "1. Environment Setup: ", "Create and activate virtual environment:\n`python -m venv venv`\n`venv\\Scripts\\activate` (Windows) or `source venv/bin/activate` (Linux/macOS)")
    add_bullet_p(doc, "2. Dependencies: ", "Install required Python packages:\n`pip install -r backend/requirements.txt`")
    add_bullet_p(doc, "3. Database Migrations: ", "Execute schema migrations:\n`python manage.py makemigrations`\n`python manage.py migrate`")
    add_bullet_p(doc, "4. Create Superuser: ", "Create initial administrative HOD account:\n`python manage.py createsuperuser`")
    add_bullet_p(doc, "5. Run Local Server: ", "Launch Django development server:\n`python manage.py runserver` -> Access at http://localhost:8000")
    add_bullet_p(doc, "6. Run Pytest Verification: ", "Execute full Django automated unit test suite:\n`$env:PYTHONPATH='backend'`\n`venv\\Scripts\\python -m pytest backend/ --ds=college_management_system.test_settings`")
    
    # Save Document
    output_path = os.path.join(os.getcwd(), "College_ERP_Comprehensive_Project_Analysis.docx")
    doc.save(output_path)
    print(f"DOCUMENT_CREATED_SUCCESSFULLY: {output_path}")

if __name__ == "__main__":
    main()
